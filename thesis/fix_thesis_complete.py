"""
논문 종합 수정 스크립트
1. 본문 중복 장절항 번호 수정
2. 목차 재구성 (제 3 장 추가, 페이지 번호 업데이트)
3. 표목차 교체 (실제 본문 표로)
4. 그림목차 정리
"""

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
import copy
import re
import shutil

DOC_PATH = r'C:\Users\USER\Desktop\안티그래비티\thesis\thesis.docx.docx'
OUTPUT   = r'C:\Users\USER\Desktop\안티그래비티\thesis\thesis_final.docx'

# Word 페이지 → 논문 페이지 변환 (오프셋 16)
OFFSET = 16
def tp(word_page):
    """thesis page"""
    if word_page > OFFSET:
        return str(word_page - OFFSET)
    else:
        # 전반부: 로마자
        roman = {1:'ⅰ',2:'ⅱ',3:'ⅲ',4:'ⅳ',5:'ⅴ',6:'ⅵ',7:'ⅶ',8:'ⅷ',
                 9:'ⅸ',10:'ⅹ',11:'ⅺ',12:'ⅻ',13:'ⅹⅲ',14:'ⅹⅳ',15:'ⅹⅴ',16:'ⅹⅵ'}
        return roman.get(word_page, str(word_page))

# ── 실제 페이지 번호 (win32com으로 수집) ──────────────────────
PAGE = {
    # 전반부
    '표 목 차':        9,   # Word page
    '그림목차':        12,
    '감사의 글':       13,
    '논문개요':        15,
    # 본문 (Word pages)
    '제 1 장':         17,
    '제 1 절 연구 배경': 17,
    '제 2 절 연구 목적': 20,
    '제 3 절 연구 방법': 22,
    '제 4 절 온톨로지 핵심': 25,   # (구 제 3 절 → 제 4 절로 수정)
    '제 2 장':         27,
    '제 1 절 건설 도메인': 27,
    '제 1 항 규칙':     28,
    '제 2 항 온톨로지 기반': 29,
    '제 3 항 딥러닝':   30,
    '제 4 항 LLM':      31,
    '제 5 항 도시하수도 분야': 32,
    '제 2 절 온톨로지(Ontology)': 33,
    '제 3 절 인간 인지-행동': 35,  # 추정 (직접 검색 필요)
    '제 1 항 휴먼에러':  37,
    '제 2 항 4단계 메커니즘': 38,
    '제 3 항 위험감수성 개념': 39,
    '제 4 항 위험감수성 온톨로지': 41,  # (구 제 3 항 → 제 4 항으로 수정)
    '제 4 절 선행연구':  43,
    '제 1 항 시맨틱 웹': 43,
    '제 2 항 본 연구 차별성': 49,
    '제 3 장':          54,
    '제 1 절 연구대상 공사': 54,
    '제 2 절 도메인 지식': 57,
    '제 3 절 안전 도메인 문서': 66,
    '제 1 항 안전 도메인 온톨로지': 66,
    '제 2 항 법 체계':   72,
    '제 4 절 위험요인 도출': 83,
    '제 5 절 근로자 관리자': 92,
    '제 6 절 안전보건 온톨로지 설계': 96,
    '제 1 항 행동 메커니즘': 96,
    '제 2 항 온톨로지 클래스': 99,
    '제 3 항 RDF/OWL':   105,
    '제 4 장':           110,
    '제 1 절 상황 맥락':  110,
    '제 2 절 SWRL':      112,
    '제 1 항 법적 규정':  112,
    '제 3 절 지능형 HSE': 114,
    '제 1 항 실시간 위험지수': 114,
    '제 2 항 행동 단계별': 117,
    '제 5 장':            118,
    '제 1 절 가상 시나리오': 118,
    '제 2 절 추론 결과':   120,
    '제 3 절 시스템 처리': 120,  # 추정
    '제 6 장':            122,
    '제 1 절 연구 요약':   123,
    '제 2 절 학술적':      124,
    '제 3 절 한계점':      125,  # 추정
    '참고문헌':            127,  # 추정
    '부    록':            138,  # 추정
    'Abstract':            140,  # 추정
}


def set_para_text(para, new_text):
    """단락의 텍스트를 유지하면서 내용만 교체 (서식 첫 run 유지)"""
    # 기존 runs 보관 (서식 복사용)
    runs = para.runs
    if not runs:
        para.text = new_text
        return

    # 첫 번째 run 서식 유지하고 텍스트 변경
    first_run = runs[0]

    # 탭 기준 분리
    if '\t' in new_text:
        tab_idx = new_text.rfind('\t')
        title_part = new_text[:tab_idx]
        page_part = new_text[tab_idx:]  # \t포함
    else:
        title_part = new_text
        page_part = ''

    # 모든 run 제거
    for run in runs:
        run._element.getparent().remove(run._element)

    # 새 run 추가
    new_run = para.add_run(title_part)
    new_run.font.bold = first_run.font.bold
    new_run.font.size = first_run.font.size

    if page_part:
        page_run = para.add_run(page_part)
        page_run.font.bold = first_run.font.bold
        page_run.font.size = first_run.font.size


def update_tab_page(para, new_page_str):
    """목차 단락의 탭 뒤 페이지 번호만 교체"""
    full = para.text
    if '\t' not in full:
        return
    tab_pos = full.rfind('\t')
    old_page = full[tab_pos+1:].strip()
    if old_page == new_page_str:
        return

    # XML에서 직접 마지막 run의 텍스트 수정
    p_elem = para._element
    runs = p_elem.findall(f'.//{qn("w:r")}')

    # 마지막 run이 페이지 번호
    if runs:
        last_run = runs[-1]
        t_elem = last_run.find(qn('w:t'))
        if t_elem is not None:
            old_t = t_elem.text or ''
            # 공백 보존
            leading = re.match(r'^(\s*)', old_t).group(1)
            t_elem.text = leading + new_page_str
            if new_page_str.strip():
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def insert_para_after(ref_para, new_text, style_name='바탕글'):
    """ref_para 뒤에 새 단락 삽입"""
    ref_p = ref_para._element
    new_p = copy.deepcopy(ref_p)

    # 텍스트 교체
    for r in new_p.findall(f'.//{qn("w:r")}'):
        new_p.remove(r)

    # 새 run 생성
    new_r = etree.SubElement(new_p, qn('w:r'))
    new_t = etree.SubElement(new_r, qn('w:t'))
    new_t.text = new_text
    if ' ' in new_text:
        new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    ref_p.addnext(new_p)
    return new_p


# ═══════════════════════════════════════════════════
print('[1] 문서 로드 및 백업...')
shutil.copy2(DOC_PATH, DOC_PATH.replace('.docx', '_before_final.docx'))
doc = Document(DOC_PATH)
paras = doc.paragraphs
print(f'    단락 수: {len(paras)}')


# ═══════════════════════════════════════════════════
print('[2] 본문 중복 장절 번호 수정...')

BODY_FIXES = {
    # (단락 인덱스, 현재 텍스트 포함 키워드, 변경 내용)
    # 1장: 두 번째 제 3 절 → 제 4 절
    195: ('온톨로지 기반 건설안전', '제 3 절', '제 4 절'),
    # 2장 3절: 두 번째 제 3 항 → 제 4 항
    283: ('위험감수성의 온톨로지', '제 3 항', '제 4 항'),
    # 3장: 중복 제 2 절 → 제 3 절
    409: ('안전 도메인 문서', '제 2 절', '제 3 절'),
    # 3장: 세 개 중복 제 3 절
    491: ('공정별 위험요인', '제 3 절', '제 4 절'),
    526: ('근로자', '제 3 절', '제 5 절'),
    545: ('안전보건 오토롤지', '제 3 절', '제 6 절'),
    # 3장 6절 하위: 중복 제 2 항 → 제 3 항
    602: ('RDF/OWL', '제 2 항', '제 3 항'),
}

for idx, (keyword, old_prefix, new_prefix) in BODY_FIXES.items():
    if idx >= len(paras):
        continue
    para = paras[idx]
    text = para.text
    if keyword in text and old_prefix in text and '\t' not in text:
        new_text = text.replace(old_prefix, new_prefix, 1)
        for run in para.runs:
            if old_prefix in run.text:
                run.text = run.text.replace(old_prefix, new_prefix, 1)
                print(f'    본문 [{idx}]: "{old_prefix}" → "{new_prefix}" | {text[:50]}')
                break
    else:
        print(f'    !! 본문 [{idx}] 매칭 실패: "{text[:60]}"')


# ═══════════════════════════════════════════════════
print('[3] 목차 수정...')

# 현재 TOC 구조 (python-docx 인덱스 기준)
# 목차 항목 텍스트 변경 및 페이지 업데이트

TOC_TEXT_FIXES = {
    # idx: (필수 포함 키워드, 새 전체 텍스트)
    # 1장 항목들
    13: ('연구 범위',       '제 3 절 연구 방법론적 체계 및 범위\t' + tp(22)),
    14: ('온톨로지 핵심',   '제 4 절 온톨로지 기반 건설안전 지식모델의 핵심 용어 체계\t' + tp(25)),
    # 2장: 항 번호 수정
    22: ('하수도 공사 분야 법규', '제 5 항 도시 하수도 공사 분야 법규 자동 검토 연구의 공백\t' + tp(32)),
    # 2장 3절 내 항 번호 수정
    28: ('위험감수성의 온톨로지', '제 4 항  위험감수성의 온톨로지 스키마 연계\t' + tp(41)),
    # 2장 → Chapter 3 구조 변경
    32: ('도시하수도 공사 안전 보건  지식', '제 3 장 도시하수도 공사 안전보건 온톨로지 및 지식 인프라 설계\t' + tp(54)),
    33: ('연구 대상의 공사개요',    '제 1 절 연구대상의 공사개요\t' + tp(54)),
    34: ('하수도 공사 도메인 지식 정의', '제 2 절 하수도 공사 도메인 지식 정의 및 범위 선정\t' + tp(57)),
    35: ('안전 도메인 문서',         '제 3 절 도시하수도 공사 안전 도메인 문서\t' + tp(66)),
    36: ('안전 도메인 온톨로지 구축', '  제 1 항 도시하수도 공사 안전 도메인 온톨로지 구축\t' + tp(66)),
    37: ('산업안전보건법',           '  제 2 항 법 체계에 따른 건설안전 관리 업무 비교 분석\t' + tp(72)),
    38: ('도시 하수도 위험요인 도출',  '제 4 절 도시 하수도 공정별 위험요인 도출\t' + tp(83)),
    40: ('도시하수도 공정별 위험 요인', ''),  # 제 4 절로 통합, 이 항목 제거
    41: ('근로자·관리자',             '제 5 절 근로자·관리자 행동 메커니즘 기반 온톨로지 모델링\t' + tp(92)),
    42: ('메커니즘 정의',             '  제 1 항 행동 메커니즘 정의 및 보건관리 연계 스키마 설계\t' + tp(96)),
    43: ('Protégé',                   '  제 2 항 온톨로지 클래스(Class) 계층 구조 설계\t' + tp(99)),
    44: ('RDF/OWL',                   '  제 3 항 RDF/OWL 기반 기계 판독형 지식베이스 구축 및 SPARQL 검증\t' + tp(105)),
    # 4장 항목들
    46: ('실시간 공정위험',           '제 4 장 실시간 위험추론 및 안전 모니터링 시스템 개발\t' + tp(110)),
    47: ('상황 맥락',                 '제 1 절 상황 맥락(Context) 데이터 수집 및 시맨틱 매핑\t' + tp(110)),
    48: ('SWRL',                      '제 2 절 SWRL 규칙 기반 실시간 안전보건 위험 추론 엔진 설계\t' + tp(112)),
    49: ('법적 규정 이행도',           '  제 1 항 법적 규정 이행도 중심의 정량화 알고리즘\t' + tp(112)),
    50: ('공정별 안전지수',            ''),  # 본문에 없음, 제거
    51: ('지능형 HSE',                '제 3 절 지능형 HSE 안전보건 관제 대시보드 구현\t' + tp(114)),
    52: ('실시간 위험지수',            '  제 1 항 실시간 위험지수 시각화 및 Trigger 알림 체계\t' + tp(114)),
    53: ('행동 메커니즘 단계별',       '  제 2 항 행동 메커니즘 단계별 데이터 축적 및 적응형 학습 환경 구현\t' + tp(117)),
    # 5장
    55: ('시스템 구현 결과',           '제 5 장 시스템 구현 결과 및 타당성 검증\t' + tp(118)),
    56: ('가상 도시하수도',            '제 1 절 가상 도시하수도 공사 시나리오 기반 실험 설계\t' + tp(118)),
    57: ('안전보건 위험 추론 결과',    '제 2 절 안전보건 위험 추론 결과 및 대시보드 구동 평가\t' + tp(120)),
    58: ('시스템 처리',               '제 3 절 시스템 처리 성능 및 실무적 타당성 분석\t' + tp(120)),
    # 6장
    60: ('결론',                      '제 6 장 결론\t' + tp(122)),
    61: ('연구의 요약',               '제 1 절 연구의 요약 및 최종 성과\t' + tp(123)),
    62: ('학술적 독창성',              '제 2 절 학술적 독창성 및 실무적 기대효과\t' + tp(124)),
    63: ('한계점',                    '제 3 절 연구의 한계점 및 향후 연구과제\t' + tp(125)),
    # 기타
    65: ('참고문헌',                  '참고문헌\t' + tp(127)),
    66: ('부',                        '부    록\t' + tp(138)),
    67: ('Abstract',                  'Abstract\t' + tp(140)),
}

# 페이지 번호도 업데이트가 필요한 항목 (텍스트 변경 없이)
TOC_PAGE_ONLY = {
    10: tp(17),   # 제 1 장 서론
    11: tp(17),   # 제 1 절 연구 배경
    12: tp(4),    # 제 2 절 연구 목적
    16: tp(11),   # 제 2 장
    17: tp(11),   # 제 1 절 건설 도메인
    18: tp(12),   # 제 1 항 규칙
    19: tp(13),   # 제 2 항 온톨로지
    20: tp(14),   # 제 3 항 자연어처리
    21: tp(15),   # 제 4 항 LLM
    23: tp(17),   # 제 2 절 온톨로지
    24: tp(19),   # 제 3 절 인간
    25: tp(21),   # 제 1 항 휴먼에러
    26: tp(22),   # 제 2 항 4단계
    27: tp(23),   # 제 3 항 위험감수성 개념
    29: tp(27),   # 제 4 절 선행연구
    30: tp(27),   # 제 1 항 시맨틱
    31: tp(33),   # 제 2 항 본 연구
    # 전반부
    5:  tp(9),    # 표 목 차
    6:  tp(12),   # 그림목차
    7:  tp(13),   # 감사의 글
    8:  tp(15),   # 논문개요
}

# TOC 텍스트 변경 적용
for idx, (keyword, new_text) in TOC_TEXT_FIXES.items():
    if idx >= len(paras):
        print(f'    !! 인덱스 초과: {idx}')
        continue
    para = paras[idx]
    old_text = para.text

    if not old_text.strip():
        continue

    if keyword and keyword not in old_text:
        print(f'    !! [{idx}] 키워드 "{keyword}" 미발견 in: "{old_text[:60]}"')
        continue

    if new_text == '':
        # 항목 제거 (빈 텍스트로 만들기)
        for run in para.runs:
            run.text = ''
        print(f'    [{idx}] 항목 제거: "{old_text[:60]}"')
    else:
        # 전체 텍스트 교체
        tab_pos = new_text.rfind('\t')
        new_title = new_text[:tab_pos] if tab_pos >= 0 else new_text
        new_page = new_text[tab_pos+1:] if tab_pos >= 0 else ''

        if para.runs:
            # title 부분을 첫 run에 설정
            # 마지막 run에 페이지 설정
            runs = para.runs
            if len(runs) >= 2:
                # 탭 포함 runs 구조 파악
                combined = ''.join(r.text for r in runs)
                if '\t' in combined:
                    # 탭 위치까지가 제목, 이후가 페이지
                    tab_in_combined = combined.rfind('\t')
                    # runs 전체를 새 내용으로 재설정
                    remaining = new_title
                    for i, run in enumerate(runs):
                        if i < len(runs) - 1:
                            run.text = remaining[:len(run.text)] if remaining else ''
                        else:
                            run.text = '\t' + new_page
                        remaining = remaining[len(run.text):]
                    # 단순하게: 첫 run에 제목, 마지막 run에 페이지
                    runs[0].text = new_title
                    for r in runs[1:-1]:
                        r.text = ''
                    runs[-1].text = '\t' + new_page
                else:
                    runs[0].text = new_title + '\t' + new_page
            else:
                runs[0].text = new_title + '\t' + new_page

        print(f'    [{idx}] → "{new_text[:80]}"')

# 페이지 번호만 업데이트
for idx, new_page in TOC_PAGE_ONLY.items():
    if idx >= len(paras):
        continue
    update_tab_page(paras[idx], new_page)


# ═══════════════════════════════════════════════════
print('[4] 표목차 교체 (실제 본문 표로)...')

# 실제 본문 표 캡션과 페이지 (win32com에서 확인한 값)
REAL_TABLES = [
    ('<표 2-1> 온톨로지 선행연구 분석',               tp(27)),
    ('<표 2-2> 선행연구의 핵심 접근 방법 및 기술적 격차 분석', tp(33)),
    ('<표 3-1> 산업안전보건법과 건설기술진흥법 비교',     tp(56)),
    ('<표 3-2> 공정 별 위험 요인 분석 및 도출',         tp(67)),
    ('<표 3-3> 공정별 안전지수(SSI) 8대 영역과 온톨로지 클래스 대응 체계', tp(80)),
    ('<표 4-1> CRI 등급별 Trigger 알림 및 관리자 조치 체계', tp(98)),
    ('<표 5-1> 가상 도시하수도 공사 시나리오 기반 설계',   tp(118)),
    ('<표 5-2> 시나리오 별 추론 결과',                   tp(120)),
]

# 표목차 영역: 단락 71~116
# 처음 8개는 실제 내용으로 채우고, 나머지는 비움
TABLE_TOC_START = 71
TABLE_TOC_END = 116

for i, (title, page) in enumerate(REAL_TABLES):
    idx = TABLE_TOC_START + i
    if idx >= len(paras):
        break
    para = paras[idx]
    new_text = f'{title}\t{page}'
    if para.runs:
        para.runs[0].text = title
        for r in para.runs[1:-1]:
            r.text = ''
        if len(para.runs) > 1:
            para.runs[-1].text = f'\t{page}'
        else:
            para.runs[0].text = new_text
    print(f'    표목차 [{idx}]: {new_text}')

# 남은 항목 비우기
for idx in range(TABLE_TOC_START + len(REAL_TABLES), TABLE_TOC_END + 1):
    if idx >= len(paras):
        break
    para = paras[idx]
    for run in para.runs:
        run.text = ''


# ═══════════════════════════════════════════════════
print('[5] 그림목차 정리...')

# 본문에 그림 캡션이 없으므로 기존 잘못된 항목 제거
# 향후 작성자가 직접 추가 필요
FIGURE_TOC_START = 120
FIGURE_TOC_END = 138

for idx in range(FIGURE_TOC_START, FIGURE_TOC_END + 1):
    if idx >= len(paras):
        break
    para = paras[idx]
    text = para.text.strip()
    if text and '\t' in text and ('<그림' in text or '<Figure' in text):
        for run in para.runs:
            run.text = ''
        print(f'    그림목차 [{idx}] 제거: "{text[:60]}"')


# ═══════════════════════════════════════════════════
print('[6] 저장...')
doc.save(OUTPUT)
print(f'    저장 완료: {OUTPUT}')
print()
print('=== 완료 ===')
print('수동 확인 필요 사항:')
print('  1. 3장 절/항 번호가 본문에서 올바르게 수정되었는지 확인')
print('  2. 그림목차: 본문에 그림 캡션(<그림 N-N> 형식)을 추가한 후 목차에 기입 필요')
print('  3. 참고문헌/부록/Abstract 페이지 번호 정확성 확인 필요')
print(f'  4. Word에서 열어 레이아웃 확인 후 최종 검토 권장')
