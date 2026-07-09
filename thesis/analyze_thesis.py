"""
thesis.docx.docx 구조 완전 분석 스크립트
"""
import re
from docx import Document

DOCX_PATH = r"C:\Users\USER\Desktop\안티그래비티\thesis\thesis.docx.docx"

doc = Document(DOCX_PATH)

print("=" * 80)
print(f"총 단락 수 (python-docx): {len(doc.paragraphs)}")
print(f"총 표 수: {len(doc.tables)}")
print("=" * 80)

# ============================
# 1. TOC 표 내용 확인
# ============================
print("\n" + "=" * 80)
print("[ 1. TOC 표 내용 ]")
print("=" * 80)

toc_labels = {3: "목차 (tables[3])", 4: "표목차 (tables[4])", 5: "그림목차 (tables[5])"}

for tbl_idx, label in toc_labels.items():
    if tbl_idx >= len(doc.tables):
        print(f"\n  tables[{tbl_idx}] 존재하지 않음")
        continue
    tbl = doc.tables[tbl_idx]
    print(f"\n--- {label} ---")
    print(f"  행 수: {len(tbl.rows)}, 열 수: {len(tbl.columns) if tbl.rows else 0}")
    for r_i, row in enumerate(tbl.rows):
        for c_i, cell in enumerate(row.cells):
            text = cell.text.strip()
            if text:
                print(f"  [행{r_i:3d}|열{c_i}] {repr(text)}")


# ============================
# 2. 본문 실제 챕터 구조 (단락 145 이후)
# ============================
print("\n" + "=" * 80)
print("[ 2. 본문 실제 챕터 구조 (단락 145 이후) ]")
print("=" * 80)

# 패턴: "제 N 장", "제 N 절", "제 N 항" 포함 (탭 없음)
chapter_pattern = re.compile(r'^제\s*\d+\s*[장절항]')

print(f"\n  ※ '제 N 장/절/항' 패턴 탐색 (인덱스 145 이후):\n")

chapter_entries = []
for i, para in enumerate(doc.paragraphs):
    if i < 145:
        continue
    text = para.text.strip()
    if not text:
        continue
    # 탭 없는 것 (탭 문자 제외)
    if '\t' in para.text:
        continue
    if chapter_pattern.match(text):
        style = para.style.name if para.style else "?"
        chapter_entries.append((i, text, style))
        print(f"  [{i:5d}] ({style}) {text}")

print(f"\n  총 챕터/절/항 수: {len(chapter_entries)}")


# ============================
# 3. 본문 실제 표 캡션
# ============================
print("\n" + "=" * 80)
print("[ 3. 본문 실제 표 캡션 (단락 145 이후) ]")
print("=" * 80)

# <표...> 또는 표 N-N 등 패턴
table_caption_pattern = re.compile(r'(?:<표|표\s*[\d가-힣]|〈표|【표|\[표)', re.IGNORECASE)

table_captions = []
for i, para in enumerate(doc.paragraphs):
    if i < 145:
        continue
    text = para.text.strip()
    if not text:
        continue
    if '\t' in para.text:
        continue
    if table_caption_pattern.search(text):
        style = para.style.name if para.style else "?"
        table_captions.append((i, text, style))

print(f"\n  표 캡션 단락 ({len(table_captions)}개):\n")
for i, text, style in table_captions:
    print(f"  [{i:5d}] ({style}) {text}")


# ============================
# 4. 본문 실제 그림 캡션
# ============================
print("\n" + "=" * 80)
print("[ 4. 본문 실제 그림 캡션 (단락 145 이후) ]")
print("=" * 80)

fig_caption_pattern = re.compile(r'(?:<그림|그림\s*[\d가-힣]|〈그림|Figure|Fig\.?\s*\d|\[그림)', re.IGNORECASE)

fig_captions = []
for i, para in enumerate(doc.paragraphs):
    if i < 145:
        continue
    text = para.text.strip()
    if not text:
        continue
    if '\t' in para.text:
        continue
    if fig_caption_pattern.search(text):
        style = para.style.name if para.style else "?"
        fig_captions.append((i, text, style))

print(f"\n  그림 캡션 단락 ({len(fig_captions)}개):\n")
for i, text, style in fig_captions:
    print(f"  [{i:5d}] ({style}) {text}")


# ============================
# 5. TOC 목차 vs 본문 헤딩 불일치
# ============================
print("\n" + "=" * 80)
print("[ 5. TOC 목차 vs 본문 헤딩 불일치 분석 ]")
print("=" * 80)

# TOC에서 챕터 항목 추출
toc_entries = []
if 3 < len(doc.tables):
    tbl = doc.tables[3]
    for row in tbl.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text and chapter_pattern.match(text):
                toc_entries.append(text)

print(f"\n  TOC에서 추출된 챕터 항목 ({len(toc_entries)}개):")
for t in toc_entries:
    print(f"    {t}")

print(f"\n  본문에서 찾은 챕터 항목 ({len(chapter_entries)}개):")
for i, text, style in chapter_entries:
    print(f"    [{i}] {text}")

# 정규화 함수: 공백 제거하여 비교
def normalize(s):
    return re.sub(r'\s+', '', s)

toc_set = {normalize(t): t for t in toc_entries}
body_set = {normalize(text): (i, text) for i, text, style in chapter_entries}

print("\n  --- 목차에는 있지만 본문에 없는 항목 ---")
missing_in_body = [orig for norm, orig in toc_set.items() if norm not in body_set]
if missing_in_body:
    for t in missing_in_body:
        print(f"    [목차만] {t}")
else:
    print("    (없음)")

print("\n  --- 본문에는 있지만 목차에 없는 항목 ---")
missing_in_toc = [(i, text) for norm, (i, text) in body_set.items() if norm not in toc_set]
if missing_in_toc:
    for i, text in missing_in_toc:
        print(f"    [본문만] [{i}] {text}")
else:
    print("    (없음)")

# 장/절/항 번호 기반으로 제목 불일치 비교
def extract_number(s):
    """제 N 장/절/항 에서 숫자와 유형 추출"""
    m = re.match(r'^제\s*(\d+)\s*([장절항])\s*(.*)', s)
    if m:
        return (m.group(1), m.group(2), m.group(3).strip())
    return None

toc_numbered = {}
for t in toc_entries:
    parsed = extract_number(t)
    if parsed:
        key = (parsed[0], parsed[1])  # (번호, 장/절/항)
        toc_numbered[key] = parsed[2]  # 제목

body_numbered = {}
for i, text, style in chapter_entries:
    parsed = extract_number(text)
    if parsed:
        key = (parsed[0], parsed[1])
        body_numbered[key] = (i, parsed[2])

print("\n  --- 장/절/항 번호는 같지만 제목이 다른 항목 ---")
mismatched = []
for key in toc_numbered:
    if key in body_numbered:
        toc_title = toc_numbered[key]
        body_title = body_numbered[key][1]
        body_idx = body_numbered[key][0]
        if normalize(toc_title) != normalize(body_title):
            mismatched.append((key, toc_title, body_idx, body_title))

if mismatched:
    for (num, typ), toc_t, b_idx, body_t in mismatched:
        print(f"    제 {num} {typ} | 목차: '{toc_t}' | 본문[{b_idx}]: '{body_t}'")
else:
    print("    (없음)")


# ============================
# 추가: 표목차/그림목차 내용 vs 본문 캡션 비교
# ============================
print("\n" + "=" * 80)
print("[ 추가: 표목차(tables[4]) 내용 vs 본문 표 캡션 비교 ]")
print("=" * 80)

if 4 < len(doc.tables):
    tbl4 = doc.tables[4]
    toc_table_captions = []
    for row in tbl4.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text and table_caption_pattern.search(text):
                toc_table_captions.append(text)

    print(f"\n  표목차에 있는 항목 ({len(toc_table_captions)}개):")
    for t in toc_table_captions:
        print(f"    목차: {t}")

    print(f"\n  본문 표 캡션 ({len(table_captions)}개):")
    for i, text, style in table_captions:
        print(f"    본문[{i}]: {text}")

print("\n" + "=" * 80)
print("[ 추가: 그림목차(tables[5]) 내용 vs 본문 그림 캡션 비교 ]")
print("=" * 80)

if 5 < len(doc.tables):
    tbl5 = doc.tables[5]
    toc_fig_captions = []
    for row in tbl5.rows:
        for cell in row.cells:
            text = cell.text.strip()
            if text and fig_caption_pattern.search(text):
                toc_fig_captions.append(text)

    print(f"\n  그림목차에 있는 항목 ({len(toc_fig_captions)}개):")
    for t in toc_fig_captions:
        print(f"    목차: {t}")

    print(f"\n  본문 그림 캡션 ({len(fig_captions)}개):")
    for i, text, style in fig_captions:
        print(f"    본문[{i}]: {text}")

print("\n" + "=" * 80)
print("분석 완료")
print("=" * 80)
